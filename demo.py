import streamlit as st
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext
from llama_index.llms import OpenAI
from PyPDF2 import PdfReader
import io
from dotenv import load_dotenv
import tempfile
from llama_index import Document
from llama_index.query_engine.retriever_query_engine import RetrieverQueryEngine
from llama_index.indices.vector_store.retrievers import VectorIndexRetriever
from docx import Document as DocxDocument
from io import BytesIO
import markdown as markdown
import openai
load_dotenv()

# Set up Streamlit app
st.set_page_config(layout='wide')
st.title('DEMO')
input_column, response_column = st.columns([2,3])
uploaded_file = input_column.file_uploader("Choose a PDF file", type="pdf")

formality = input_column.slider('Formality', 0, 100, 50)

# Add inputs for user context


# Add inputs for user goal
userprompt = input_column.text_input('Prompt')

# Add dropdown for response structure


language = input_column.selectbox('Language', ['English','', 'Spanish', 'French', 'German', 'Italian', 'Turkish'])

query_templates = {
    "Response": f"I need a detailed and comprehensive response for the given document. {userprompt}. Please provide the response in markdown format with appropriate formatting and styles.",
}

selected_description = input_column.selectbox('Select a query template', list(query_templates.keys()))

# Then, when you're ready to generate the query:
selected_template = query_templates[selected_description]

# Add slider for temperature
temperature = input_column.slider('Temperature', 0.0, 1.0, 0.5)

# Add slider for max tokens
max_tokens = input_column.slider('Max Tokens', 100, 2000, 500)

if uploaded_file is not None:
    # Create the ServiceContext with the user-selected temperature
    service_context = ServiceContext.from_defaults(llm=OpenAI(temperature=temperature, model="gpt-4", max_tokens=max_tokens))

    with st.spinner('Reading PDF...'):
        pdf = PdfReader(io.BytesIO(uploaded_file.getvalue()))
        text = " ".join(page.extract_text() for page in pdf.pages)
        documents = [Document(text=text)]
        index = VectorStoreIndex.from_documents(documents, service_context=service_context)

    if input_column.button('Generate'):
        # Determine formality phrase
        if formality < 33:
            formality = "In a casual and conversational style, "
        elif formality < 67:
            formality = "In a neutral style, "
        else:
            formality = "In a highly formal and academic style, "

        # Add user context and structure to the query
        query = selected_template.format(userprompt=userprompt, language=language, formality=formality)
        # Generate the response
# Generate the response
        with st.spinner(f'Generating Response...'):
            retriever = VectorIndexRetriever(index=index)
            query_engine = RetrieverQueryEngine(retriever=retriever)
            response = query_engine.query(query)
            # Store the response text in the session state
            st.session_state['response'] = response.response

    # Display the response stored in the session state
    if 'response' in st.session_state:
        openai.api_key = st.secrets["OPENAI_API_KEY"]
        response_text = st.session_state['response']
        response_column.markdown(response_text)
        # Create a new Document
        doc = DocxDocument()
        doc.add_paragraph(response_text)
        # Save the document to a BytesIO buffer
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        # Add a download button for the response
        response_column.download_button("Download DOCX", buf.getvalue(), file_name="response.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


input_column.markdown("<p style='text-align: center;'> Demo </p>", unsafe_allow_html=True)