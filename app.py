import streamlit as st
from llama_index import VectorStoreIndex, SimpleDirectoryReader, ServiceContext
from llama_index.llms import OpenAI
from PyPDF2 import PdfReader
import io
from dotenv import load_dotenv
import tempfile
from llama_index import Document
from llama_index.query_engine.retriever_query_engine import RetrieverQueryEngine
from llama_index.indices.vector_store.retrievers import VectorIndexRetriever
import os
import openai
import pyrebase
import markdown
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from io import BytesIO

st.set_page_config(layout='wide')
firebase_config = {
    "apiKey": st.secrets["firebase"]["apiKey"],
    "authDomain": st.secrets["firebase"]["authDomain"],
    "databaseURL": st.secrets["firebase"]["databaseURL"],
    "projectId": st.secrets["firebase"]["projectId"],
    "storageBucket": st.secrets["firebase"]["storageBucket"],
    "messagingSenderId": st.secrets["firebase"]["messagingSenderId"],
    "appId": st.secrets["firebase"]["appId"]
}

firebase = pyrebase.initialize_app(firebase_config)
auth = firebase.auth()

# Initialize session state
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False

# Create login form
if not st.session_state['logged_in']:
    username = st.text_input('Username')
    password = st.text_input('Password', type='password')

    if st.button('Log In'):
        try:
            user = auth.sign_in_with_email_and_password(username, password)
            st.session_state['logged_in'] = True
            
        except:
            st.error('Invalid username/password')

if st.session_state['logged_in']:
    # Retrieve OpenAI API key from secrets
    openai.api_key = st.secrets["OPENAI_API_KEY"]

    # Set up Streamlit app
    st.image("onepager-logo-tr.png")
    input_column, response_column = st.columns([2,3])

    # Add inputs for sender, recipient, and purpose
    sender = input_column.text_input('Who sends this OnePager?')
    recipient = input_column.text_input('Who receives the OnePager?')
    purpose = input_column.text_input('What is the purpose of the OnePager?')

    # Add dropdown for document structure
    doc_structure = input_column.radio('How should the OnePager be structured?', ['AI Suggestion', 'Bullet Points', 'Pitch (3 Parts)', 'Report', 'No Structure'], horizontal=True)
    # Add sliders for tone, technicality, and length

    formality_labels = {1: 'Casual', 2: 'Somewhat Casual', 3: 'Neutral', 4: 'Somewhat Formal', 5: 'Formal'}
    tone_value = input_column.slider('How formal should the OnePager be?', 1, 5, 3)
    tone = formality_labels[tone_value]
    input_column.write('Selected formality: ' + tone)
    technicality = input_column.slider('How technical should the OnePager be formulated?', 1, 5, 3, format="%d")
    length = input_column.slider('How detailed should the OnePager be?', 1, 5, 3, format="%d")

    # Add file uploader for background information
    uploaded_file = input_column.file_uploader("Upload a PDF with background information.", type="pdf")

    # Add inputs for source description, call to action, and additional info
    source_description = input_column.text_input('What kind of document is this? Why is it relevant?')
    has_call_to_action = input_column.checkbox('Should the OnePager give a recommendation for action?')
    call_to_action = input_column.text_input('What is the recommendation for action?')
    action_tone = input_column.slider('How directly should this recommendation be placed?', 1, 5, 3, format="%d")
    additional_info = input_column.text_input('What additional information belongs in the OnePager?')

    # Add slider for temperature
    temperature = input_column.slider('Temperature', 0.0, 1.0, 0.5)

    # Add slider for max tokens
    max_tokens = input_column.slider('Max Tokens', 100, 2000, 500)

    if uploaded_file is not None:
        # Create the ServiceContext with the user-selected temperature
        service_context = ServiceContext.from_defaults(llm=OpenAI(temperature=temperature, model="gpt-4", max_tokens=max_tokens))

        with st.spinner('Reading PDF...'):
            pdf = PdfReader(io.BytesIO(uploaded_file.getvalue()))
            text = " ".join(page.extract_text() for page in pdf.pages)
            documents = [Document(text=text)]
            index = VectorStoreIndex.from_documents(documents, service_context=service_context)

        if input_column.button('Generate'):
            # Determine formality phrase
            if tone_value == 1:
                formality = "In a casual and conversational style, "
            elif tone_value == 2:
                formality = "In a somewhat casual style, "
            elif tone_value == 3:
                formality = "In a neutral style, "
            elif tone_value == 4:
                formality = "In a somewhat formal style, "
            else:
                formality = "In a highly formal and academic style, "

            # Determine action tone phrase
            if action_tone == 1:
                action_formality = "In a casual and conversational style, "
            elif action_tone == 2:
                action_formality = "In a somewhat casual style, "
            elif action_tone == 3:
                action_formality = "In a neutral style, "
            elif action_tone == 4:
                action_formality = "In a somewhat formal style, "
            else:
                action_formality = "In a highly formal and academic style, "

            # Add user context and structure to the query
            if doc_structure == 'AI Suggestion':
                query = f"As {sender}, I need a document for {recipient} that is {length} in length and {technicality} in technicality. My goal is {purpose}. I want the response in English. Please provide the response in markdown format with appropriate features. {formality}"
            else:
                query = f"As {sender}, I need a {doc_structure} of the document for {recipient} that is {length} in length and {technicality} in technicality. My goal is {purpose}. I want the response in English. Please provide the response in markdown format with appropriate features. {formality}"
            
            # If source_description is provided, add it to the query
            if source_description:
                query += f" The source document is: {source_description}."

            # If has_call_to_action is selected, add call to action to the query
            if has_call_to_action:
                query += f" Generate a full and extensive {doc_structure} with a recommendation for action: {call_to_action}. {action_formality} The additional information is: {additional_info}."
            else:
                query += f" Generate a full and extensive {doc_structure}."

            # Generate the response
            with st.spinner(f'Generating {doc_structure.lower()}...'):
                retriever = VectorIndexRetriever(index=index)
                query_engine = RetrieverQueryEngine(retriever=retriever)
                response = query_engine.query(query)
                # Store the response text in the session state
                st.session_state['response'] = response.response

    # Display the response stored in the session state
    if 'response' in st.session_state:
        response_text = st.session_state['response']
        response_column.markdown(response_text)
        # Convert markdown to HTML
        html = markdown.markdown(response_text)

        # Parse HTML
        soup = BeautifulSoup(html, 'html.parser')

        # Create a new Document
        doc = DocxDocument()

        # Add each paragraph to the document
        for element in soup:
            if element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                doc.add_heading(element.text, level=int(element.name[1]))
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListBullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='ListNumber')
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        # Add a download button for the response
        response_column.download_button("Download response", buf.getvalue(), file_name="response.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")